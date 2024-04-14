import os
import re
from django.shortcuts import render, redirect
from django.http import HttpResponse
import pdfplumber
from .forms import FileForm
from docx import Document
import pandas as pd
import aspose.words as aw

def upload_file(request):

    def detect_file_format(file_path):
        _, extension = os.path.splitext(file_path)
        if extension.lower() == '.pdf':
            return 'pdf'
        elif extension.lower() == '.docx':
            return 'docx'
        elif extension.lower() == '.doc':
            return 'doc'
        else:
            return 'Unknown'
    
    def extract_text_from_pdf(pdf_path):
        with pdfplumber.open(pdf_path) as pdf:
            text = ''
            for page in pdf.pages:
                text += page.extract_text()
        return text
    
    def extract_text_from_docx(docx_path):
        doc = Document(docx_path)
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + ' '
        return text
    
    def extract_emails(text):
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails
        
    def extract_contact_numbers(text):
        phone_pattern = r'\b(?:\+\d{1,2}\s)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b'
        phone_numbers = re.findall(phone_pattern, text)
        return phone_numbers
    

    if request.method == 'POST':
        form = FileForm(request.POST, request.FILES)
        if form.is_valid():
            data = []
            for file in request.FILES.getlist('file'):
                form = FileForm({'name': file.name}, {'file': file})
                filePath = 'media/uploads/' + file.name

                if detect_file_format(file.name) == 'pdf':
                    if form.is_valid():
                        form.save()
                    text = extract_text_from_pdf(filePath)
                    emails = extract_emails(text)
                    ph_no = extract_contact_numbers(text)
                    data.append({'File Name': file.name, 'Email': emails, 'Contact Number': ph_no, 'Text': text})
                    
                elif detect_file_format(file.name) == 'docx':
                    if form.is_valid():
                        form.save()
                    text = extract_text_from_docx(filePath)
                    emails = extract_emails(text)
                    ph_no = extract_contact_numbers(text)
                    data.append({'File Name': file.name, 'Email': emails, 'Contact Number': ph_no, 'Text': text})

                elif detect_file_format(file.name) == 'doc':
                    if form.is_valid():
                        form.save()
                    doc = aw.Document(filePath)
                    docx_path = os.path.splitext(filePath)[0] + '.docx'
                    doc.save(docx_path)
                    text = extract_text_from_docx(docx_path)
                    emails = extract_emails(text)
                    ph_no = extract_contact_numbers(text)
                    data.append({'File Name': file.name, 'Email': emails, 'Contact Number': ph_no, 'Text': text})
            
            df = pd.DataFrame(data)

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Disposition'] = 'attachment; filename=output.xlsx'

            with pd.ExcelWriter(response, engine='xlsxwriter') as writer:
                df.to_excel(writer, index=False)

            return response
                
    else:
        form = FileForm()
    return render(request, 'index.html', {'form': form})